# -*- coding: utf-8 -*-
"""
掃描「學生作業」下各學生資料夾，找出 final HTML、評分表 CSV、雷達圖 SVG、
文字理念、形容詞、靈感紙圖片、AI對話內容連結，產出 gallery 資料。
執行：python build_gallery_data.py
結論三檔從 Data/ 讀取；gallery_data.json、gallery_embedded.js 寫入 Gallery/。
掃描時排除：Script、Data、Doc、Gallery、__pycache__。
"""
import csv
import json
import os
import re

try:
    from paths_config import BASE_DIR, DATA_DIR, GALLERY_DIR, iter_student_folders
except ImportError:
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    DATA_DIR = os.path.join(BASE_DIR, "Data")
    GALLERY_DIR = os.path.join(BASE_DIR, "Gallery")
    def iter_student_folders():
        skip = {"Script", "Data", "Doc", "Gallery", "__pycache__"}
        for name in sorted(os.listdir(BASE_DIR)):
            path = os.path.join(BASE_DIR, name)
            if os.path.isdir(path) and name not in skip:
                yield name


def find_final_html(folder_path):
    """在資料夾與子資料夾中找檔名含 final 的 .html/.htm，回傳第一個的相對路徑（相對於 base）。"""
    folder_path = os.path.normpath(folder_path)
    base = os.path.dirname(folder_path)
    for root, _dirs, files in os.walk(folder_path):
        for f in sorted(files):
            name_lower = f.lower()
            if "final" not in name_lower:
                continue
            if name_lower.endswith(".html") or name_lower.endswith(".htm"):
                abs_path = os.path.join(root, f)
                return os.path.relpath(abs_path, base)
    return None


def parse_folder_name(folder_name):
    """資料夾名為 姓名_情緒，回傳 (姓名, 情緒)。若無底線則情緒為空。"""
    if "_" in folder_name:
        parts = folder_name.split("_", 1)
        return parts[0], parts[1]
    return folder_name, ""


def parse_avg_stars(score_rows):
    """從評分表 CSV 的得分欄（★ 數量）計算平均星等，回傳 1.0~5.0 或 None。"""
    if not score_rows or len(score_rows) < 2:
        return None
    total = 0
    count = 0
    for row in score_rows[1:]:
        if len(row) >= 3:
            stars = row[2].count("★")
            if 0 <= stars <= 5:
                total += stars
                count += 1
    if not count:
        return None
    return round(total / count, 1)


def find_file_in_folder(folder_path, base_dir, filename):
    """在資料夾與子資料夾中找檔名完全符合的檔案，回傳相對路徑或 None。"""
    folder_path = os.path.normpath(folder_path)
    for root, _dirs, files in os.walk(folder_path):
        for f in files:
            if f == filename:
                abs_path = os.path.join(root, f)
                return os.path.relpath(abs_path, base_dir).replace("\\", "/")
    return None


def read_text_file(abs_path, encoding="utf-8-sig"):
    """讀取文字檔內容，失敗回傳 None。"""
    if not abs_path or not os.path.isfile(abs_path):
        return None
    try:
        with open(abs_path, "r", encoding=encoding) as f:
            return f.read().strip()
    except Exception:
        try:
            with open(abs_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception:
            return None


# 匹配 [連結 1] https://... 或 [連結 2] https://...（整份檔案都會掃，故可抓到多個連結）
AI_CHAT_LINK_PATTERN = re.compile(r"\[連結\s*\d+\]\s*(https?://[^\s\r\n]+)")


def extract_ai_chat_links_from_text(text):
    """從字串中擷取所有 [連結 N] 後的 URL，回傳 list of str。"""
    if not text:
        return []
    return AI_CHAT_LINK_PATTERN.findall(text)


def extract_ai_chat_links(abs_path):
    """從 AI對話內容.txt 擷取所有 [連結 N] 後的 URL，回傳 list of str。"""
    if not abs_path or not os.path.isfile(abs_path):
        return []
    try:
        with open(abs_path, "r", encoding="utf-8-sig") as f:
            text = f.read()
    except Exception:
        return []
    return extract_ai_chat_links_from_text(text)


def extract_ai_chat_links_from_docx(abs_path):
    """從 .docx 抽出文字後擷取 [連結 N] URL，回傳 list of str。"""
    if not abs_path or not os.path.isfile(abs_path):
        return []
    try:
        from docx import Document
        doc = Document(abs_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return extract_ai_chat_links_from_text(text)
    except Exception:
        return []


def extract_ai_used_from_score_csv(score_rows):
    """從評分表 CSV 的「使用 AI」欄位擷取不重複值，回傳 list（如 ["Perplexity", "Gemini"]）。"""
    if not score_rows or len(score_rows) < 2:
        return []
    header = score_rows[0]
    col_idx = None
    for i, cell in enumerate(header):
        c = (cell or "").strip().replace(" ", "")
        if c == "使用AI" or (cell and "使用" in cell and "AI" in cell):
            col_idx = i
            break
    if col_idx is None:
        return []
    seen = set()
    result = []
    for row in score_rows[1:]:
        if len(row) <= col_idx:
            continue
        val = (row[col_idx] or "").strip()
        if val and val not in seen:
            seen.add(val)
            result.append(val)
    return result


# 從 AI 對話連結網址推斷使用的 AI，回傳不重複的 list（如 ["ChatGPT", "Gemini"]）
# 辨識順序：先比對具體網域，再比對關鍵字，盡量減少「其他」
def infer_ai_from_urls(urls):
    if not urls:
        return []
    ai_map = []
    for url in urls:
        url_lower = url.lower()
        if "chatgpt.com" in url_lower or "openai.com" in url_lower:
            ai_map.append("ChatGPT")
        elif "claude.ai" in url_lower or "anthropic.com" in url_lower:
            ai_map.append("Claude")
        elif (
            "gemini.google.com" in url_lower
            or "aistudio.google.com" in url_lower  # Google AI Studio 分享連結
            or ("g.co" in url_lower and "gemini" in url_lower)  # 短網址 https://g.co/gemini/share/xxx
            or ("google.com" in url_lower and "gemini" in url_lower)
        ):
            ai_map.append("Gemini")
        elif "copilot" in url_lower or "bing.com" in url_lower:
            ai_map.append("Copilot")
        elif "perplexity.ai" in url_lower:
            ai_map.append("Perplexity")
        elif "meta.ai" in url_lower or ("facebook.com" in url_lower and "/ai" in url_lower):
            ai_map.append("Meta AI")
        elif "grok.x.ai" in url_lower or ("x.ai" in url_lower and "grok" in url_lower) or ("x.com" in url_lower and "grok" in url_lower):
            ai_map.append("Grok")
        elif "deepseek.com" in url_lower:
            ai_map.append("DeepSeek")
        elif "moonshot.cn" in url_lower or "kimi" in url_lower:
            ai_map.append("Kimi")
        elif "mistral.ai" in url_lower:
            ai_map.append("Mistral")
        elif "coze.com" in url_lower:
            ai_map.append("Coze")
        elif "doubao.com" in url_lower:
            ai_map.append("豆包")
        elif "yiyan.baidu.com" in url_lower or "wenxin.baidu.com" in url_lower:
            ai_map.append("文心一言")
        elif "qianwen.aliyun.com" in url_lower or "dashscope.aliyun.com" in url_lower:
            ai_map.append("通義千問")
        elif "01.ai" in url_lower or "yi.com" in url_lower:
            ai_map.append("零一萬物")
        else:
            ai_map.append("其他")
    # 去重、保留順序
    seen = set()
    result = []
    for x in ai_map:
        if x not in seen:
            seen.add(x)
            result.append(x)
    return result


def find_inspiration_images(folder_path, base_dir):
    """
    在資料夾與子資料夾中找圖片檔（靈感紙等），檔名含「靈感」的優先。
    支援：靈感紙、圖像靈感、文字靈感 等類似命名，副檔名含 .heic、.dng。
    回傳相對路徑 list。
    """
    folder_path = os.path.normpath(folder_path)
    image_ext = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".heic", ".dng")
    # 優先：檔名含「靈感」者（含 圖像靈感.heic、文字靈感.heic、靈感紙.jpg 等）
    with_lingan = []
    other = []
    for root, _dirs, files in os.walk(folder_path):
        for f in files:
            low = f.lower()
            if not any(low.endswith(ext) for ext in image_ext):
                continue
            abs_path = os.path.join(root, f)
            rel = os.path.relpath(abs_path, base_dir).replace("\\", "/")
            if "靈感" in f:
                with_lingan.append(rel)
            else:
                other.append(rel)
    return with_lingan if with_lingan else other


# 文字理念 / 形容詞 支援的副檔名，皆嘗試抽出文字直接顯示（需安裝對應套件）
# 建議：pip install python-docx pypdf striprtf odfpy
CONCEPT_EXTS = (".txt", ".doc", ".docx", ".pdf", ".rtf", ".odt")


def extract_text_from_file(abs_path, ext):
    """
    依副檔名從檔案抽出純文字。.txt 直接讀取；.docx/.pdf/.rtf/.odt 需對應套件。
    回傳抽出之文字或 None。
    """
    if not abs_path or not os.path.isfile(abs_path):
        return None
    ext = ext.lower()
    if ext == ".txt":
        return read_text_file(abs_path)
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(abs_path)
            return "\n".join(p.text for p in doc.paragraphs).strip() or None
        except Exception:
            return None
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(abs_path)
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            return "\n".join(parts).strip() or None
        except Exception:
            return None
    if ext == ".rtf":
        try:
            from striprtf import striprtf
            with open(abs_path, "r", encoding="utf-8", errors="replace") as f:
                return striprtf.stript_rtf(f.read()).strip() or None
        except Exception:
            return None
    if ext == ".odt":
        try:
            from odf import text, teletype
            from odf.opendocument import load
            doc = load(abs_path)
            parts = []
            for el in doc.getElementsByType(text.P):
                parts.append(teletype.extractText(el))
            return "\n".join(parts).strip() or None
        except Exception:
            return None
    # .doc 舊版 Word 二進位格式，一般需另用工具，此處不處理
    return None


def find_concept_file(folder_path, base_dir, base_names, allow_stem_endswith=None, allow_stem_contains=None):
    """
    在資料夾與子資料夾中找檔名符合條件的檔案並抽出文字。
    條件：stem 在 base_names 內、或 endswith allow_stem_endswith、或 stem 含 allow_stem_contains 任一字串。
    檔名去副檔名後 strip 再比對。回傳 (內容或 None)。
    """
    if isinstance(base_names, str):
        base_names = [base_names]
    folder_path = os.path.normpath(folder_path)
    found_files = []
    for root, _dirs, files in os.walk(folder_path):
        for f in files:
            stem, ext = os.path.splitext(f)
            stem = stem.strip()
            ext = ext.lower()
            if ext not in CONCEPT_EXTS:
                continue
            ok = stem in base_names
            if not ok and allow_stem_endswith:
                ok = any(stem.endswith(s) for s in allow_stem_endswith)
            if not ok and allow_stem_contains:
                ok = any(s in stem for s in allow_stem_contains)
            if not ok:
                continue
            abs_path = os.path.join(root, f)
            rel = os.path.relpath(abs_path, base_dir).replace("\\", "/")
            found_files.append((os.path.join(base_dir, rel), ext))
    if not found_files:
        return None
    found_files.sort(key=lambda x: (0 if x[1] == ".txt" else 1, x[0]))
    for abs_path, ext in found_files:
        content = extract_text_from_file(abs_path, ext)
        if content:
            return content
    return None


# 形容詞若為圖片（如 感受轉視覺.jpg）則存相對路徑，藝廊可顯示
ADJECTIVES_IMAGE_STEMS = ("感受轉視覺",)


def find_adjectives_image(folder_path, base_dir):
    """找形容詞類圖片（檔名去副檔名後在 ADJECTIVES_IMAGE_STEMS 內）。回傳相對路徑或 None。"""
    folder_path = os.path.normpath(folder_path)
    image_ext = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".heic", ".dng")
    for root, _dirs, files in os.walk(folder_path):
        for f in files:
            stem, ext = os.path.splitext(f)
            stem = stem.strip()
            if ext.lower() not in image_ext:
                continue
            if stem in ADJECTIVES_IMAGE_STEMS:
                abs_path = os.path.join(root, f)
                return os.path.relpath(abs_path, base_dir).replace("\\", "/")
    return None


def main():
    base_dir = BASE_DIR
    entries = []
    emotions_set = set()

    # 讀取 prompt 分析與同事評語（步驟 3 產出，可選）
    prompt_feedback = {}
    feedback_path = os.path.join(DATA_DIR, "prompt_feedback.json")
    if os.path.isfile(feedback_path):
        try:
            with open(feedback_path, "r", encoding="utf-8") as f:
                prompt_feedback = json.load(f)
        except Exception:
            pass

    for name in iter_student_folders():
        folder = os.path.join(base_dir, name)
        # 略過非「姓名_情緒」格式的資料夾
        if "_" not in name:
            continue
        student_name, emotion = parse_folder_name(name)
        final_html_rel = find_final_html(folder)
        if not final_html_rel:
            continue
        score_csv_name = f"{name}_評分表.csv"
        radar_svg_name = f"{name}_雷達圖.svg"
        score_csv_path = os.path.join(folder, score_csv_name)
        radar_svg_path = os.path.join(folder, radar_svg_name)
        score_rows = []
        if os.path.isfile(score_csv_path):
            try:
                with open(score_csv_path, "r", encoding="utf-8-sig") as f:
                    reader = csv.reader(f)
                    score_rows = list(reader)
            except Exception:
                pass
        radar_svg_content = ""
        if os.path.isfile(radar_svg_path):
            try:
                with open(radar_svg_path, "r", encoding="utf-8") as f:
                    radar_svg_content = f.read()
            except Exception:
                pass
        if emotion:
            emotions_set.add(emotion)
        avg_stars = parse_avg_stars(score_rows)

        # 理念：文字理念、設計理念、文字敘述，或檔名含「文字理念」/ 以「設計理念」結尾
        text_concept = find_concept_file(
            folder,
            base_dir,
            ["文字理念", "設計理念", "文字敘述"],
            allow_stem_endswith=("設計理念",),
            allow_stem_contains=("文字理念",),
        )
        # 形容詞：文字檔 形容詞／一句話形容／一句形容；圖片檔 感受轉視覺.jpg 等
        adjectives = find_concept_file(folder, base_dir, ["形容詞", "一句話形容", "一句形容"])
        adjectives_image_rel = find_adjectives_image(folder, base_dir)

        # AI對話內容連結：從 .txt 擷取所有 [連結 N] URL；再從 AI對話串網址.docx 補上（合併、去重）
        ai_chat_links = []
        for cand in ["AI對話內容.txt", "AI對話內容(1).txt"]:
            p = os.path.join(folder, cand)
            if os.path.isfile(p):
                ai_chat_links = extract_ai_chat_links(p)
                break
        if not ai_chat_links:
            for root, _dirs, files in os.walk(folder):
                for f in files:
                    if f.endswith(".txt") and "AI對話內容" in f:
                        ai_chat_links = extract_ai_chat_links(os.path.join(root, f))
                        if ai_chat_links:
                            break
                if ai_chat_links:
                    break
        # 若有 AI對話串網址.docx，一併擷取連結並合併（保留順序、去重）
        for root, _dirs, files in os.walk(folder):
            for f in files:
                if f.endswith(".docx") and "AI對話串網址" in f:
                    extra = extract_ai_chat_links_from_docx(os.path.join(root, f))
                    seen = set(ai_chat_links)
                    for url in extra:
                        if url not in seen:
                            seen.add(url)
                            ai_chat_links.append(url)
                    break
            else:
                continue
            break
        # 若該生交的是 AI 對話 PDF（如 AI對話內容.pdf），加入相對路徑供藝廊點擊開啟
        for root, _dirs, files in os.walk(folder):
            for f in files:
                if f.endswith(".pdf") and "AI對話" in f:
                    abs_pdf = os.path.join(root, f)
                    rel_pdf = os.path.relpath(abs_pdf, base_dir).replace("\\", "/")
                    if rel_pdf not in ai_chat_links:
                        ai_chat_links.append(rel_pdf)
            break

        inspiration_images = find_inspiration_images(folder, base_dir)
        ai_used = infer_ai_from_urls(ai_chat_links)
        # 若連結推不出來（空）或只推出「其他」（例如僅有 PDF 路徑），改以評分表「使用 AI」欄位
        csv_ai = extract_ai_used_from_score_csv(score_rows)
        if not ai_used or (len(ai_used) == 1 and ai_used[0] == "其他"):
            if csv_ai:
                ai_used = csv_ai

        entry = {
            "folderName": name,
            "studentName": student_name,
            "emotion": emotion,
            "avgStars": avg_stars,
            "finalHtmlRel": final_html_rel.replace("\\", "/"),
            "scoreCsvRel": f"{name}/{score_csv_name}".replace("\\", "/") if os.path.isfile(score_csv_path) else None,
            "radarSvgRel": f"{name}/{radar_svg_name}".replace("\\", "/") if os.path.isfile(radar_svg_path) else None,
            "scoreRows": score_rows,
            "radarSvgContent": radar_svg_content if radar_svg_content else None,
            "textConcept": text_concept,
            "adjectives": adjectives,
            "adjectivesImageRel": adjectives_image_rel,
            "inspirationImageRels": inspiration_images,
            "aiChatLinks": ai_chat_links,
            "aiUsed": ai_used,
        }
        # 併入 Data/prompt_feedback.json（analyze_prompt_and_feedback.py 產出）
        fb = prompt_feedback.get(name, {})
        if fb:
            entry["promptCount"] = fb.get("promptCount", 0)
            entry["tokenCount"] = fb.get("tokenCount", 0)
            entry["colleagueComment"] = fb.get("colleagueComment", "")
            entry["prompts"] = fb.get("prompts") or []
        entries.append(entry)

    ais_set = set()
    for e in entries:
        used = e.get("aiUsed") or []
        if used:
            for ai in used:
                ais_set.add(ai)
        else:
            ais_set.add("其他")
    ais_set.add("其他")
    data = {
        "entries": entries,
        "emotions": sorted(emotions_set),
        "ais": sorted(ais_set),
    }

    # 結論三檔從 Data/ 讀取
    conclusion_report_path = os.path.join(DATA_DIR, "評分總覽報告.txt")
    conclusion_per_student_path = os.path.join(DATA_DIR, "評分總覽_各生.csv")
    conclusion_chart_path = os.path.join(DATA_DIR, "評分總覽_圖表資料.csv")
    data["conclusionReport"] = read_text_file(conclusion_report_path) or ""
    data["conclusionPerStudentCsv"] = read_text_file(conclusion_per_student_path) or ""
    data["conclusionChartDataCsv"] = read_text_file(conclusion_chart_path) or ""

    os.makedirs(GALLERY_DIR, exist_ok=True)
    json_path = os.path.join(GALLERY_DIR, "gallery_data.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"已寫入 {json_path}，共 {len(entries)} 位學生。")

    js_path = os.path.join(GALLERY_DIR, "gallery_embedded.js")
    with open(js_path, "w", encoding="utf-8") as f:
        f.write("window.GALLERY_DATA = ")
        f.write(json.dumps(data, ensure_ascii=False))
        f.write(";\n")
    print(f"已寫入 {js_path}")


if __name__ == "__main__":
    main()
